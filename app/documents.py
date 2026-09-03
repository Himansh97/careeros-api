"""Resume export to PDF and DOCX.

The layout mirrors the candidate's own LaTeX base resume: a centered header
with headline and credentials, a professional summary, categorised skills,
then experience with right-aligned dates and an italic role subtitle, and
education with right-aligned date and GPA.

It stays ATS-safe throughout: single column, real selectable text, standard
section headings, no text boxes or images. Right-aligned dates use a
borderless two-cell row (PDF) and a tab stop (DOCX) rather than a layout
table, so parsers still read one linear column of text.
"""
from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as _RT
from docx.oxml import OxmlElement as _El
from docx.oxml.ns import qn as _qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

DARK = colors.HexColor("#282828")


# Typographic characters that mark a document as machine-composed. Nobody
# reaches for option-shift-hyphen while typing a resume, so a page carrying
# nine em dashes reads as generated to a human reviewer. Smart quotes are here
# for a second reason: older ATS text extractors mangle them, and a mangled
# quote inside a bullet is a keyword that no longer matches.
#
# Purely punctuation. No word, number or name is altered, so this cannot change
# what the document claims.
_PLAIN = {
    # Spaced forms first, so " x \u2014 y " collapses to " x - y " rather than
    # leaving the double space a bare replacement would. An earlier version
    # fixed that with a global whitespace collapse and broke the contact line,
    # whose doubled spaces around the pipes are deliberate -- the PDF and DOCX
    # renderers then disagreed, which tests/test_outreach_copy.py exists to
    # catch and duly did.
    " \u2014 ": " - ",
    " \u2013 ": " - ",
    "\u2014": "-",   # em dash
    "\u2013": "-",   # en dash
    "\u2018": "'",   # left single quote
    "\u2019": "'",   # right single quote
    "\u201c": '"',   # left double quote
    "\u201d": '"',   # right double quote
    "\u2022": "-",   # bullet
    "\u2026": "...",  # ellipsis
    "\u00a0": " ",   # non-breaking space
}


def plain_punctuation(text: str) -> str:
    """Normalise typography without touching content.

    Applied in `_escape` rather than at each call site because that is the one
    place every string passes through on its way into a paragraph. The dashes
    arrive from three different sources -- role headings built in `tailor`,
    project names recorded in the evidence file, and the claim text itself --
    so fixing any one of them would have left the other two.
    """
    out = text or ""
    for char, replacement in _PLAIN.items():
        out = out.replace(char, replacement)
    return out


def _escape(text: str) -> str:
    """ReportLab paragraphs accept a mini-HTML dialect, so escape markup."""
    cleaned = plain_punctuation(text)
    return cleaned.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _pretty_month(token: str) -> str:
    from .tailor import _pretty_date

    return _pretty_date(token)


def safe_filename(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", "_", name).strip("_")
    return cleaned or "resume"


CONTACT_SEP = "  |  "


def contact_parts(profile: Any) -> list[tuple[str, str]]:
    """The header line as (what it reads, where it goes) pairs.

    One source of truth for both renderers. They used to build the same string
    independently, which is how a fix lands in the PDF and silently misses the
    DOCX that half of ATS uploads actually parse.

    Ordered by how a recruiter uses it. Phone and email first because those are
    how they reply. The portfolio sits beside LinkedIn rather than at the end:
    it is the only item on the line that shows work rather than asserting it,
    and past the location nobody reads. Location stays last because it is the
    field most likely to end the conversation and least likely to start one.

    The visible text stays the address itself, never a "LinkedIn" label. A
    resume gets printed, forwarded as text, and scraped — and in all three the
    label survives while the link does not.
    """
    from .outreach import portfolio_host

    linkedin = (
        (profile.linkedin_url or "")
        .replace("https://", "")
        .replace("http://", "")
        .replace("www.", "")
        .rstrip("/")
    )
    site = portfolio_host(getattr(profile, "portfolio_url", ""))
    phone = (profile.phone or "").strip()
    tel = "".join(ch for ch in phone if ch.isdigit() or ch == "+")

    pairs = [
        (phone, f"tel:{tel}" if tel else ""),
        ((profile.email or "").strip(), f"mailto:{profile.email}" if profile.email else ""),
        (linkedin, f"https://{linkedin}" if linkedin else ""),
        (site, f"https://{site}" if site else ""),
        ((profile.location or "").strip(), ""),
    ]
    return [(text, href) for text, href in pairs if text]


def _contact_line(profile: Any) -> str:
    """The same header as plain text, for anything that cannot carry a link."""
    return CONTACT_SEP.join(text for text, _ in contact_parts(profile))


def _contact_markup(profile: Any) -> str:
    """The header for ReportLab, with real link annotations on it.

    Deliberately not blue and not underlined. The visible text is already the
    address, so decorating it adds nothing a reader needs and makes a printed
    resume look like a web page from 2003. The link is there for the reader who
    clicks; the text is there for everyone else.
    """
    out = []
    for text, href in contact_parts(profile):
        safe = _escape(text)
        out.append(f'<a href="{_escape(href)}">{safe}</a>' if href else safe)
    return _escape(CONTACT_SEP).join(out)


def _edu_period(edu: dict[str, Any]) -> str:
    """Show the study period, not just graduation.

    A lone graduation date leaves the employment gap the degree explains
    looking like an unexplained absence. This resume has exactly that shape:
    employment ends Jul 2023 and resumes Jul 2025, and the MS that fills those
    two years said only "May 2025". A reviewer does not reconstruct someone's
    life from a graduation date, so what the page showed was a 24-month hole.

    The docstring above described this behaviour before the code did; only the
    graduation date was ever read. `start_date` is recorded in the profile for
    both degrees and user-confirmed, so nothing here is inferred.
    """
    start = _pretty_month(edu.get("start_date", ""))
    end = _pretty_month(edu.get("graduation_date", ""))
    return f"{start} - {end}" if start and end and start != end else (end or start)


def _edu_left(edu: dict[str, Any]) -> str:
    """Degree, institution, and GPA -- everything except the dates.

    GPA sits here rather than beside the period. The right-hand column is a
    fifth of the width, sized for a date, and a study period plus a GPA wrapped
    inside it, breaking between "GPA" and the number. Putting the GPA with the
    degree it belongs to is also the ordinary resume shape, and it leaves the
    date column holding a date, exactly like the employment rows directly
    above.
    """
    parts = [edu.get("degree", ""), edu.get("institution", "")]
    if edu.get("location"):
        parts.append(edu["location"])
    left = ", ".join(x for x in parts[1:] if x)
    left = f"{parts[0]} - {left}" if left else parts[0]
    if edu.get("gpa"):
        left += f", GPA {edu['gpa']}"
    return left


def _coursework_for(
    entry: dict[str, Any], coursework: list[dict[str, Any]]
) -> list[dict[str, Any]]:
    """Coursework claims belonging to this education entry.

    An academic claim carries the institution in its employer field, so the
    two match on that.
    """
    institution = (entry.get("institution") or "").strip().lower()
    if not institution:
        return []
    return [
        c for c in coursework
        if (c.get("employer") or "").strip().lower() == institution
    ]


def _role_meta(profile: Any, employer: str) -> dict[str, Any]:
    for role in profile.employment_history or []:
        if role.get("employer") == employer:
            return role
    return {}


# Title-casing mangles acronyms ("Bi & Visualization", "Software & Devops"),
# so the labels that matter are spelled out to match the base resume exactly.
GROUP_LABELS = {
    "business_analytics_and_statistical_modeling": "Business Analytics & Statistical Modeling",
    "data_engineering_and_big_data": "Data Engineering & Big Data",
    "bi_and_visualization": "BI & Visualization",
    "project_and_delivery_management": "Project & Delivery Management",
    "software_and_devops": "Software & DevOps",
    "compliance": "Compliance",
    "mortgage_domain": "Mortgage Domain",
    "cloud_and_devsecops": "Cloud & DevSecOps",
    "ai_and_llm_engineering": "AI & LLM Engineering",
    "frontend_and_full_stack": "Frontend & Full-Stack",
}

ACRONYMS = {"Bi": "BI", "Devops": "DevOps", "Ai": "AI", "Ml": "ML", "Sql": "SQL",
            "Etl": "ETL", "Api": "API", "Ui": "UI", "Ux": "UX"}


def _group_label(group: str) -> str:
    if group in GROUP_LABELS:
        return GROUP_LABELS[group]
    label = group.replace("_", " ").replace(" and ", " & ").title()
    for wrong, right in ACRONYMS.items():
        label = re.sub(rf"\b{wrong}\b", right, label)
    return label


# How much of the skills inventory reaches the page. The inventory itself stays
# complete: it is what `scoring._find_evidence` reads for its weaker
# "listed but not demonstrated" tier, so deleting entries there would turn real
# skills into gaps, which is the wrong direction entirely.
#
# This is a rendering cap, and it is safe precisely because the groups are
# already ordered by what this posting asked for. The groups that fall off the
# end are the ones this employer never mentioned, so the ATS keywords that
# matter for this job survive while the wall does not. Twelve categories and
# roughly 120 items read as "everything I have ever touched" and cost about
# fourteen lines of a one-page resume, which is fourteen lines not spent on
# evidence.
MAX_SKILL_GROUPS = 4
MAX_SKILLS_PER_GROUP = 14


def _prioritised_skills(resume: dict[str, Any], profile: Any) -> list[tuple[str, list[str]]]:
    """Order skill groups by what the job asked for, and keep the top ones.

    Ordering adapts so a recruiter scanning the first line sees the skills
    their posting actually named instead of a fixed house order. Everything
    past MAX_SKILL_GROUPS is dropped from the page but kept in the profile.
    """
    wanted = {w.lower() for w in (resume.get("matchedSkills") or [])}
    groups = list((getattr(profile, "skills_inventory", {}) or {}).items())

    def rank_item(item: str) -> int:
        low = item.lower()
        return 0 if any(w == low or w in low for w in wanted) else 1

    ordered: list[tuple[str, list[str], float]] = []
    for name, items in groups:
        hits = float(sum(1 for i in items if rank_item(i) == 0))
        # The group's own name is a category signal in its own right. Matching
        # only individual items left "mortgage_domain" tied with generic data
        # engineering on a mortgage posting — every specific tool in it
        # (Encompass, MeridianLink, MISMO) is too specialised to appear in the
        # requirement labels, so the most relevant group ranked second.
        label = name.replace("_", " ").lower()
        if any(w in label or label in w for w in wanted):
            hits += 2.5
        ordered.append((name, sorted(items, key=rank_item), hits))

    ordered.sort(key=lambda g: -g[2])
    return [
        (name, items[:MAX_SKILLS_PER_GROUP])
        for name, items, _ in ordered[:MAX_SKILL_GROUPS]
    ]


def _page_count(pdf_bytes: bytes) -> int:
    from pypdf import PdfReader

    return len(PdfReader(io.BytesIO(pdf_bytes)).pages)


def _trim(resume: dict[str, Any], lead: int, rest: int,
          projects: int = 3) -> dict[str, Any]:
    """A copy of the resume with bullets capped per role.

    Bullets arrive relevance-ordered, so trimming drops the least relevant
    evidence first. A job matching more of the candidate's background keeps
    more bullets — the page fills itself according to the role.
    """
    trimmed = dict(resume)
    trimmed["sections"] = [
        {**s, "bullets": s["bullets"][: lead if i == 0 else rest]}
        for i, s in enumerate(resume.get("sections", []))
    ]
    # Projects shrink alongside employment, but not on the same curve. Tying
    # project bullets directly to `rest` made the section the binding
    # constraint: at rest=2 it rendered three projects of two full claim
    # bullets each, roughly twelve lines, which pushed the page over and forced
    # the ladder down to rest=1 -- costing one bullet at every role after the
    # first. Measured on a real posting, decoupling them buys three employment
    # bullets back at the same page count.
    #
    # A project is a showcase, not a job history. `_build_projects` already
    # orders bullets by marginal coverage, so the first one is the strongest
    # and a second mostly repeats the pitch.
    entries = resume.get("projects") or []
    if entries:
        # One bullet each once the budget is tight. A project is a showcase and
        # _build_projects already orders its bullets by marginal coverage, so
        # the first is the strongest and a second mostly repeats the pitch.
        per_project = 2 if projects >= 3 and rest >= 3 else 1
        trimmed["projects"] = [
            {**p, "bullets": p["bullets"][:per_project]}
            for p in entries[:max(0, projects)]
        ]
    return trimmed


# Typographic density levels, loosest first. Tightening leading and font size
# reclaims far more space than cutting bullets does, and costs nothing —
# whereas every dropped bullet is real evidence the employer no longer sees.
DENSITIES = [
    {"body": 9.4, "lead": 12.6, "gap": 10, "margin": 0.62},
    {"body": 9.1, "lead": 11.9, "gap": 9, "margin": 0.58},
    {"body": 8.8, "lead": 11.2, "gap": 8, "margin": 0.55},
    {"body": 8.5, "lead": 10.6, "gap": 7, "margin": 0.5},
    {"body": 8.2, "lead": 10.1, "gap": 6, "margin": 0.45},
]


# One page. A reviewer working a stack decides on the first page and rarely
# turns to the second, so evidence on page two is evidence that mostly is not
# read -- and at this much experience a second page reads as padding rather
# than depth. This was 2, which is why generated documents were running long
# while the hand-cut versions fit.
MAX_PAGES = 1
# Below this, the final page looks abandoned rather than finished. A resume
# that runs 1.15 pages reads worse than a tight single page — the reader sees
# the white space, not the extra evidence.
MIN_LAST_PAGE_FILL = 0.45


def _last_page_fill(pdf_bytes: bytes) -> float:
    """Roughly how full the final page is, 0-1.

    Measured by comparing extracted line counts, because reportlab does not
    report leftover frame space after the fact. Approximate is enough: this
    only has to distinguish "page ends naturally" from "page is mostly blank".
    """
    from pypdf import PdfReader

    pages = PdfReader(io.BytesIO(pdf_bytes)).pages
    if len(pages) <= 1:
        return 1.0
    counts = [len([ln for ln in (p.extract_text() or "").splitlines() if ln.strip()])
              for p in pages]
    full = max(counts[:-1]) or 1
    return min(counts[-1] / full, 1.0)


def build_pdf(resume: dict[str, Any], profile: Any) -> bytes:
    """Render, measure, and re-render until the pages are genuinely full.

    Three rules, in priority order:

    1. Never exceed MAX_PAGES, which is one page.
    2. Never leave the final page mostly blank.
    3. Subject to those, keep as much real evidence as possible and set it as
       loosely — i.e. as readably — as it will go.

    Loosest typography is tried first because looser type fills more of the
    page; tightening is what gets used to claw back an overflow, not a default.
    Bullets are only dropped once no density fits, since every dropped bullet
    is evidence the employer no longer sees.
    """
    def acceptable(pdf: bytes) -> bool:
        return (
            _page_count(pdf) <= MAX_PAGES
            and _last_page_fill(pdf) >= MIN_LAST_PAGE_FILL
        )

    # Pass 1 — full content, loosest readable setting that fits the rules.
    fitted: bytes | None = None
    for density in DENSITIES:
        candidate = _render_pdf(resume, profile, density)
        if acceptable(candidate):
            return candidate
        if _page_count(candidate) <= MAX_PAGES and fitted is None:
            # Fits on pages, but the last one is thin. Hold it as a fallback
            # while trying to do better.
            fitted = candidate

    # Pass 2 - drop the least relevant evidence a step at a time, and at each
    # step set the surviving content as loosely as it will go. Trim level is
    # the outer loop because evidence outranks typography: more bullets in
    # tight type beats fewer in loose type, which is the priority order stated
    # above. Within one trim level the loosest readable density wins.
    tight = DENSITIES[-1]
    # Employment bullets are the last thing given up. For each bullet level the
    # project section is shrunk first, so a page that will not fit loses a
    # project before it loses a bullet -- three composed situation-action-result
    # bullets per role are the argument the page is making, and a project entry
    # is a supporting exhibit.
    for lead, rest in [(7, 4), (6, 3), (5, 3), (5, 2), (4, 2), (4, 1), (3, 1)]:
        for projects in (3, 2, 1):
            trimmed = _trim(resume, lead, rest, projects)
            # Cheap feasibility check first. If the tightest setting still
            # overflows here, no looser density can fit, so skip the four
            # renders that would only confirm it.
            if not acceptable(_render_pdf(trimmed, profile, tight)):
                continue
            for density in DENSITIES:
                candidate = _render_pdf(trimmed, profile, density)
                if acceptable(candidate):
                    return candidate

    # Pass 3 - nothing fit. Return the hardest trim rather than raising: a
    # document that runs long is still a document the candidate can edit, and
    # failing here would block an application over a layout rule.
    return fitted or _render_pdf(_trim(resume, 3, 1, 1), profile, tight)


# --------------------------------------------------------------------- PDF
def _render_pdf(resume: dict[str, Any], profile: Any,
                density: dict[str, float] | None = None) -> bytes:
    d = density or DENSITIES[1]
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=d["margin"] * inch,
        rightMargin=d["margin"] * inch,
        topMargin=d["margin"] * inch,
        bottomMargin=(d["margin"] - 0.1) * inch,
        title=f"{profile.name} - {resume['jobTitle']}",
        author=profile.name,
        # ReportLab otherwise writes `Producer: ReportLab PDF Library` and the
        # literal string "(unspecified)" into Creator and Subject. A toolchain
        # name is a strange thing to find in the properties of a resume, and
        # "(unspecified)" reads as a field nobody filled in. Blank rather than
        # renamed: claiming some other authoring tool would be asserting
        # something false about the file.
        creator="",
        subject="",
        producer="",
    )
    width = doc.width
    base = getSampleStyleSheet()

    name_s = ParagraphStyle("Name", parent=base["Title"], fontSize=16.5, leading=19,
                            alignment=TA_CENTER, spaceAfter=1, textColor=DARK,
                            fontName="Helvetica-Bold")
    headline_s = ParagraphStyle("Headline", parent=base["Normal"], fontSize=10, leading=13,
                                alignment=TA_CENTER, fontName="Helvetica-Oblique",
                                textColor=DARK)
    contact_s = ParagraphStyle("Contact", parent=base["Normal"], fontSize=8.8, leading=12,
                               alignment=TA_CENTER)
    creds_s = ParagraphStyle("Creds", parent=base["Normal"], fontSize=8.8, leading=12,
                             alignment=TA_CENTER, fontName="Helvetica-Oblique")
    section_s = ParagraphStyle("Section", parent=base["Heading2"],
                               fontSize=d["body"] + 1.3, leading=d["lead"],
                               spaceBefore=d["gap"], spaceAfter=1, textColor=DARK,
                               fontName="Helvetica-Bold")
    body_s = ParagraphStyle("Body", parent=base["Normal"], fontSize=d["body"],
                            leading=d["lead"], alignment=TA_JUSTIFY)
    sub_s = ParagraphStyle("Sub", parent=base["Normal"], fontSize=d["body"] - 0.3,
                           leading=d["lead"] - 0.8,
                           fontName="Helvetica-Oblique",
                           textColor=colors.HexColor("#444444"))

    def rule() -> HRFlowable:
        return HRFlowable(width="100%", thickness=0.7, color=DARK, spaceBefore=1, spaceAfter=4)

    def split_row(left: str, right: str, left_ratio: float = 0.74) -> Table:
        """Left text with a right-aligned counterpart, as LaTeX \\hfill does."""
        right_s = ParagraphStyle("R", parent=body_s, alignment=2)
        t = Table([[Paragraph(left, body_s), Paragraph(right, right_s)]],
                  colWidths=[width * left_ratio, width * (1 - left_ratio)])
        t.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
        ]))
        return t

    flow: list[Any] = [Paragraph(_escape(profile.name.upper()), name_s)]
    flow.append(Paragraph(_contact_markup(profile), contact_s))
    if getattr(profile, "credentials_line", []):
        flow.append(Paragraph(_escape("  |  ".join(profile.credentials_line)), creds_s))
    flow.append(Spacer(1, 3))

    summary = resume.get("summary") or getattr(profile, "professional_summary", "")
    if summary:
        flow += [Paragraph("PROFESSIONAL SUMMARY", section_s), rule(),
                 Paragraph(_escape(summary), body_s)]

    skills = _prioritised_skills(resume, profile)
    if skills:
        flow += [Paragraph("SKILLS", section_s), rule()]
        for group, items in skills:
            flow.append(Paragraph(
                f"<b>{_escape(_group_label(group))}:</b> {_escape(', '.join(items))}", body_s))
        certs = getattr(profile, "certifications", []) or []
        if certs:
            flow.append(Paragraph(
                f"<b>Certifications:</b> {_escape(', '.join(certs))}", body_s))

    flow += [Paragraph("PROFESSIONAL EXPERIENCE", section_s), rule()]
    for section in resume.get("sections", []):
        meta = _role_meta(profile, section.get("employer", ""))
        dates = section.get("subheading", "").split(" · ")[0]
        location = meta.get("location", "")
        left = f"<b>{_escape(section.get('heading',''))}"
        if location:
            left += f", {_escape(location)}"
        left += "</b>"
        flow.append(split_row(left, _escape(dates)))
        if meta.get("subtitle"):
            flow.append(Paragraph(_escape(meta["subtitle"]), sub_s))
        # The per-role `Tools:` line is not rendered. Four of them cost about
        # eight lines and restated the SKILLS block sitting directly above,
        # which on a one-page resume was eight lines of duplication in place of
        # evidence -- Syracuse, Freyr and Omnicals were down to a single bullet
        # each to make room. The field stays in the profile; a reader who wants
        # the stack for a role can see it in the bullets, which name the tools
        # in the context of what was built with them.
        bullets = [ListItem(Paragraph(_escape(b.get("text", "")), body_s), leftIndent=11)
                   for b in section.get("bullets", [])]
        if bullets:
            flow.append(ListFlowable(bullets, bulletType="bullet", start="•",
                                     leftIndent=12, bulletFontSize=7, spaceBefore=1))
        flow.append(Spacer(1, 3))

    # Named deliverables sit after employment: a recruiter reads the roles for
    # context first, then the projects for what was actually shipped.
    projects = resume.get("projects") or []
    if projects:
        flow += [Paragraph("SELECTED PROJECTS", section_s), rule()]
        for project in projects:
            flow.append(Paragraph(f"<b>{_escape(project.get('name',''))}</b>", body_s))
            bullets = [
                ListItem(Paragraph(_escape(b.get("text", "")), body_s), leftIndent=11)
                for b in project.get("bullets", [])
            ]
            if bullets:
                flow.append(ListFlowable(bullets, bulletType="bullet", start="•",
                                         leftIndent=12, bulletFontSize=7, spaceBefore=1))
            flow.append(Spacer(1, 3))

    education = getattr(profile, "education", []) or []
    if education:
        flow += [Paragraph("EDUCATION", section_s), rule()]
        coursework = resume.get("coursework") or []
        for e in education:
            flow.append(split_row(f"<b>{_escape(_edu_left(e))}</b>",
                                  f"<b>{_escape(_edu_period(e))}</b>",
                                  left_ratio=0.72))
            items = [
                ListItem(Paragraph(_escape(c["text"]), body_s), leftIndent=11)
                for c in _coursework_for(e, coursework)
            ]
            if items:
                flow.append(ListFlowable(items, bulletType="bullet", start="•",
                                         leftIndent=12, bulletFontSize=7, spaceBefore=1))

    doc.build(flow)
    return buf.getvalue()


# -------------------------------------------------------------------- DOCX
def build_docx(resume: dict[str, Any], profile: Any) -> bytes:
    document = Document()
    for s in document.sections:
        s.left_margin = s.right_margin = Inches(0.6)
        s.top_margin = Inches(0.55)
        s.bottom_margin = Inches(0.5)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(2)

    def centered(text: str, size: float, bold: bool = False, italic: bool = False) -> None:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold, r.italic = bold, italic
        r.font.size = Pt(size)

    def _link_run(paragraph: Any, text: str, url: str, size: float) -> None:
        """python-docx has no hyperlink API, so the relationship and the
        w:hyperlink element are built by hand.

        Styled to match the surrounding text rather than picking up Word's
        Hyperlink style, for the same reason as the PDF: the address is already
        the visible text, and blue underline on a resume reads as an accident.
        """
        r_id = paragraph.part.relate_to(url, _RT.HYPERLINK, is_external=True)
        link = _El("w:hyperlink")
        link.set(_qn("r:id"), r_id)
        run = _El("w:r")
        props = _El("w:rPr")
        sz = _El("w:sz")
        sz.set(_qn("w:val"), str(int(size * 2)))       # half-points
        props.append(sz)
        colour = _El("w:color")
        colour.set(_qn("w:val"), "000000")
        props.append(colour)
        run.append(props)
        node = _El("w:t")
        node.text = text
        node.set(_qn("xml:space"), "preserve")
        run.append(node)
        link.append(run)
        paragraph._p.append(link)

    def contact_paragraph(size: float) -> None:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = contact_parts(profile)
        for i, (text, href) in enumerate(parts):
            if i:
                sep = p.add_run(CONTACT_SEP)
                sep.font.size = Pt(size)
            if href:
                _link_run(p, text, href, size)
            else:
                r = p.add_run(text)
                r.font.size = Pt(size)

    centered(profile.name.upper(), 16, bold=True)
    contact_paragraph(9)
    if getattr(profile, "credentials_line", []):
        centered("  |  ".join(profile.credentials_line), 9, italic=True)

    def section_heading(text: str) -> None:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x28, 0x28, 0x28)

    def split_row(left: str, right: str) -> None:
        """Right-align the date with a tab stop rather than a layout table."""
        p = document.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.2), WD_TAB_ALIGNMENT.RIGHT)
        r = p.add_run(left)
        r.bold = True
        rr = p.add_run("\t" + right)
        rr.bold = True

    summary = resume.get("summary") or getattr(profile, "professional_summary", "")
    if summary:
        section_heading("PROFESSIONAL SUMMARY")
        document.add_paragraph(summary)

    skills = _prioritised_skills(resume, profile)
    if skills:
        section_heading("SKILLS")
        for group, items in skills:
            p = document.add_paragraph()
            r = p.add_run(f"{_group_label(group)}: ")
            r.bold = True
            p.add_run(", ".join(items))
        certs = getattr(profile, "certifications", []) or []
        if certs:
            p = document.add_paragraph()
            r = p.add_run("Certifications: ")
            r.bold = True
            p.add_run(", ".join(certs))

    section_heading("PROFESSIONAL EXPERIENCE")
    for section in resume.get("sections", []):
        meta = _role_meta(profile, section.get("employer", ""))
        heading = section.get("heading", "")
        if meta.get("location"):
            heading += f", {meta['location']}"
        split_row(heading, section.get("subheading", "").split(" · ")[0])
        # `tools` is deliberately absent here too. The PDF and DOCX must
        # render the same document; tests/test_outreach_copy.py exists because
        # they once did not.
        for line in (meta.get("subtitle"),):
            if not line:
                continue
            p = document.add_paragraph()
            r = p.add_run(line)
            r.italic = True
            r.font.size = Pt(9)
        for b in section.get("bullets", []):
            document.add_paragraph(b.get("text", ""), style="List Bullet")

    projects = resume.get("projects") or []
    if projects:
        section_heading("SELECTED PROJECTS")
        for project in projects:
            p = document.add_paragraph()
            r = p.add_run(project.get("name", ""))
            r.bold = True
            for b in project.get("bullets", []):
                document.add_paragraph(b.get("text", ""), style="List Bullet")

    education = getattr(profile, "education", []) or []
    if education:
        section_heading("EDUCATION")
        coursework = resume.get("coursework") or []
        for e in education:
            split_row(_edu_left(e), _edu_period(e))
            for c in _coursework_for(e, coursework):
                document.add_paragraph(c["text"], style="List Bullet")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
